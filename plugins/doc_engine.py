import os
import json
import threading
import re

try:
    import ollama
    from openpyxl import Workbook
    from docx import Document
    from fpdf import FPDF
except ImportError:
    pass

def handle_query(query, app):
    # Route the triggers
    if query.startswith("create excel "):
        topic = query.replace("create excel ", "", 1).strip()
        build_document(topic, "excel", app)
        return True
    elif query.startswith("create word "):
        topic = query.replace("create word ", "", 1).strip()
        build_document(topic, "word", app)
        return True
    elif query.startswith("create pdf "):
        topic = query.replace("create pdf ", "", 1).strip()
        build_document(topic, "pdf", app)
        return True
        
    return False

def build_document(topic, doc_type, app):
    app.safe_append_chat("SYSTEM", f"[Document Forge] Drafting {doc_type.upper()} for: '{topic}'...")
    app.speak_async(f"Drafting the {doc_type} document.")

    def build_worker():
        try:
            # 1. The Blueprint Phase (Qwen Coder)
            if doc_type == "excel":
                prompt = f"Create a 5-row spreadsheet dataset about '{topic}'. Output ONLY a valid JSON array of arrays. The first array must be the header strings, and the next 4 arrays are the data rows. Do not include markdown or extra text."
            else:
                prompt = f"Write a professional 3-section document about '{topic}'. Output ONLY a valid JSON array of objects. Each object must have a 'heading' (string) and a 'body' (string paragraph). Do not include markdown or extra text."

            res = ollama.chat(model='qwen2.5-coder:3b', messages=[{'role': 'user', 'content': prompt}])
            
            # Extract JSON cleanly
            json_match = re.search(r'\[.*\]', res['message']['content'], re.DOTALL)
            if not json_match:
                app.safe_append_chat("ERROR", "Failed to parse structural data from Qwen.")
                return
            
            data = json.loads(json_match.group(0))
            filename = f"forge_{doc_type}_{int(threading.get_ident())}"

            # 2. The Compilation Phase
            if doc_type == "excel":
                filename += ".xlsx"
                wb = Workbook()
                ws = wb.active
                ws.title = "Dataset"
                for row in data:
                    ws.append(row)
                wb.save(filename)

            elif doc_type == "word":
                filename += ".docx"
                doc = Document()
                doc.add_heading(f"Report: {topic}", 0)
                for section in data:
                    doc.add_heading(section.get('heading', 'Section'), level=1)
                    doc.add_paragraph(section.get('body', ''))
                doc.save(filename)

            elif doc_type == "pdf":
                filename += ".pdf"
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", 'B', 16)
                pdf.cell(200, 10, txt=f"Report: {topic}", ln=True, align='C')
                pdf.ln(10)
                
                for section in data:
                    pdf.set_font("Arial", 'B', 12)
                    # Safe encoding to prevent PDF font crashes
                    safe_heading = str(section.get('heading', '')).encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(200, 10, txt=safe_heading, ln=True)
                    
                    pdf.set_font("Arial", '', 11)
                    safe_body = str(section.get('body', '')).encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 7, txt=safe_body)
                    pdf.ln(5)
                pdf.output(filename)

            # 3. Delivery
            app.safe_append_chat("SYSTEM", f"File successfully forged: {filename}")
            app.speak_async(f"{doc_type} generation complete. Opening file.")
            os.startfile(filename) # Let Windows launch the default app

        except Exception as e:
            app.safe_append_chat("ERROR", f"Omni-Forge crash: {e}")

    threading.Thread(target=build_worker, daemon=True).start()